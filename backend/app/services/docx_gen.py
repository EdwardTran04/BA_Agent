import os
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

logger = logging.getLogger("ba_agent.docx_gen")

class DOCXGenerator:
    @staticmethod
    def set_cell_background(cell, fill_hex: str):
        """Sets the background color of a table cell using XML manipulation."""
        tcPr = cell._tc.get_or_add_tcPr()
        # Remove any existing shading
        for child in list(tcPr):
            if child.tag.endswith('shd'):
                tcPr.remove(child)
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    @staticmethod
    def set_table_borders(table):
        """Applies thin, clean light-gray borders to the table to replace Word's harsh defaults."""
        tblPr = table._tbl.tblPr
        # Remove existing borders if any
        for child in list(tblPr):
            if child.tag.endswith('tblBorders'):
                tblPr.remove(child)
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
            f'  <w:left w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="EAECF0"/>'
            f'  <w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

    @staticmethod
    def set_table_margins(table, top=140, bottom=140, left=180, right=180):
        """Sets default internal padding for all cells in the table (values in dxa: 20 dxa = 1 pt)."""
        tblPr = table._tbl.tblPr
        tblCellMar = OxmlElement('w:tblCellMar')
        for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tblCellMar.append(node)
        tblPr.append(tblCellMar)

    @classmethod
    def generate_docx(cls, session_id: str, screen_name: str, screen_type: str, module: str, role: str, context: str, screen_summary: str, assumptions: list, controls: list, output_dir: str, is_draft: bool = False, *args, **kwargs) -> str:
        """
        Generates a Word document in A4 Landscape featuring only:
        1. Centered UI design image at the top
        2. Control specifications table underneath
        Fonts: Times New Roman, Size 12, Line Spacing 1.3
        """
        from datetime import datetime
        image_path = kwargs.get("image_path")

        os.makedirs(output_dir, exist_ok=True)
        suffix = "_Draft" if is_draft else ""
        now_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = os.path.join(output_dir, f"Screen_Spec_{session_id}_{now_str}{suffix}.docx")
        
        doc = Document()
        
        # 1. Page Configuration: A4 Landscape, Margins: Top=2cm, Bottom=2cm, Left=3cm, Right=2cm
        sections = doc.sections
        for section in sections:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Cm(29.7)
            section.page_height = Cm(21.0)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.0)

        # Style standard Normal style to Times New Roman 12 pt
        style_normal = doc.styles['Normal']
        font_normal = style_normal.font
        font_normal.name = 'Times New Roman'
        font_normal.size = Pt(12.0)
        font_normal.color.rgb = RGBColor(0x33, 0x3F, 0x48)  # Charcoal

        # Helper to apply line spacing 1.3, space before 6pt, space after 0pt
        def apply_ba_spacing(p):
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)

        # 2. Add Centered Document Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(0)
        title_p.paragraph_format.space_after = Pt(4)
        title_run = title_p.add_run(f"BẢNG ĐẶC TẢ CHI TIẾT MÀN HÌNH: {screen_name.upper()}")
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x1A, 0x5F, 0x7A)  # Elegant Deep Teal

        # 3. Add Centered Screen Design Picture at the top
        if image_path and os.path.exists(image_path):
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_ba_spacing(img_p)
            run_img = img_p.add_run()
            # Set size to fit nicely in A4 landscape layout (max usable width 24.7cm)
            run_img.add_picture(image_path, width=Cm(18.0))
            
            # Spacer after the picture
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(12)
            spacer.paragraph_format.space_after = Pt(0)

        # 4. Component & Control Specs Table (placed directly below the image)
        # Widths in Cm: total width = 24.7 Cm
        col_widths = [Cm(1.2), Cm(4.9), Cm(3.0), Cm(3.0), Cm(4.0), Cm(8.6)]
        headers = ["STT", "Thành phần/ Control", "Kiểu dữ liệu", "Input/ Output", "Giá trị khởi tạo", "Mô tả chi tiết"]
        
        main_table = doc.add_table(rows=len(controls) + 1, cols=6)
        main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cls.set_table_borders(main_table)
        cls.set_table_margins(main_table, top=140, bottom=140, left=180, right=180) # cell padding

        # Format header row
        hdr_row = main_table.rows[0]
        trPr = hdr_row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

        for idx, text in enumerate(headers):
            cell = hdr_row.cells[idx]
            cell.width = col_widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cls.set_cell_background(cell, "1A5F7A")  # Deep Teal Header
            
            p = cell.paragraphs[0]
            apply_ba_spacing(p)
            p.paragraph_format.space_before = Pt(4)
            
            if idx in [0, 2, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # White
            run.font.size = Pt(12.0)

        # Format data rows
        for r_idx, ctrl in enumerate(controls):
            row = main_table.rows[r_idx + 1]
            
            # Keep row together
            r_trPr = row._tr.get_or_add_trPr()
            r_trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

            # Alternating zebra shading
            bg_color = "F4F8FA" if r_idx % 2 == 1 else "FFFFFF"

            ctrl_data = [
                str(ctrl.get("STT", r_idx + 1)),
                ctrl.get("control_name", ""),
                ctrl.get("data_type", ""),
                ctrl.get("io", ""),
                ctrl.get("initial_value", ""),
                ctrl.get("description", "")
            ]

            for c_idx, val in enumerate(ctrl_data):
                cell = row.cells[c_idx]
                cell.width = col_widths[c_idx]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                cls.set_cell_background(cell, bg_color)

                p = cell.paragraphs[0]
                apply_ba_spacing(p)
                
                if c_idx in [0, 2, 3]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                run = p.add_run(val)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12.0)

                # Special column highlighting
                if c_idx == 0:
                    run.font.bold = True
                elif c_idx == 1:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        doc.save(file_path)
        logger.info(f"Successfully generated clean landscape DOCX at {file_path}")
        return os.path.abspath(file_path)
